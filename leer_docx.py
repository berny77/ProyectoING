from docx import Document

try:
    ruta = "C:/Users/bnlga/Desktop/pp3.docx"
    print(f"Ruta del archivo: {ruta}")
    doc = Document(ruta)
    print("Cantidad de párrafos:", len(doc.paragraphs))
    for p in doc.paragraphs:
        print(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
except Exception as e:
    print("Error:", e)
